"""Tests for text extraction from PDF and DOCX files."""

import pytest

from src.frontend.extraction import extract_and_combine, extract_text


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a minimal PDF file for testing."""
    from io import BytesIO

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "test.pdf"
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 720, "This is page one.")
    c.showPage()
    c.drawString(72, 720, "This is page two.")
    c.showPage()
    c.save()
    pdf_path.write_bytes(buf.getvalue())
    return str(pdf_path)


@pytest.fixture
def sample_docx(tmp_path):
    """Create a minimal DOCX file for testing."""
    import docx

    docx_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(docx_path))
    return str(docx_path)


class TestExtractText:
    def test_extract_pdf(self, sample_pdf):
        text = extract_text(sample_pdf)
        assert "page one" in text
        assert "page two" in text

    def test_extract_docx(self, sample_docx):
        text = extract_text(sample_docx)
        assert "First paragraph" in text
        assert "Second paragraph" in text

    def test_unsupported_type(self, tmp_path):
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(str(txt_path))


class TestExtractAndCombine:
    def test_single_file(self, sample_docx):
        text = extract_and_combine([sample_docx])
        assert "First paragraph" in text
        assert "MULTI-DOCUMENT" not in text

    def test_multiple_files(self, sample_pdf, sample_docx):
        text = extract_and_combine([sample_pdf, sample_docx])
        assert "MULTI-DOCUMENT PACKAGE (2 documents)" in text
        assert "DOCUMENT 1 of 2:" in text
        assert "DOCUMENT 2 of 2:" in text
        assert "page one" in text
        assert "First paragraph" in text

    def test_empty_list(self):
        assert extract_and_combine([]) == ""

    def test_filenames_in_headers(self, sample_pdf, sample_docx):
        text = extract_and_combine([sample_pdf, sample_docx])
        assert "test.pdf" in text
        assert "test.docx" in text
